"""
PDF -> Word conversion microservice (FastAPI + Docling + python-docx).

Deploy on Hugging Face Spaces (Docker), Render, or Modal, then point the web app
at https://<your-host>/convert (Backend button in the PDF to WORD tool).
"""

from __future__ import annotations

import os
import tempfile
from pathlib import Path

from fastapi import FastAPI, File, Form, HTTPException, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt

app = FastAPI(title="PDF to Word microservice")

# Browsers call this directly from the web app, so CORS must be open (or set
# ALLOWED_ORIGINS to a comma-separated list of your app origins).
app.add_middleware(
    CORSMiddleware,
    allow_origins=[o for o in os.getenv("ALLOWED_ORIGINS", "*").split(",") if o],
    allow_methods=["*"],
    allow_headers=["*"],
)

LATIN_FONT = "Arial"
DEVANAGARI_FONT = "Noto Sans Devanagari"


@app.get("/health")
def health() -> dict[str, str]:
    return {"status": "ok"}


def _style_run(run, size_pt: float = 10.5, bold: bool = False) -> None:
    run.bold = bold
    run.font.size = Pt(size_pt)
    run.font.name = LATIN_FONT
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(rfonts)
    # Unicode-safe fallbacks so Devanagari (Hindi/Marathi) never renders as boxes.
    rfonts.set(qn("w:ascii"), LATIN_FONT)
    rfonts.set(qn("w:hAnsi"), LATIN_FONT)
    rfonts.set(qn("w:cs"), DEVANAGARI_FONT)


def _add_table(document: Document, rows: list[list[str]]) -> None:
    cols = max(len(r) for r in rows)
    table = document.add_table(rows=0, cols=cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True  # full page width between margins
    for row in rows:
        cells = table.add_row().cells
        for index in range(cols):
            text = row[index] if index < len(row) else ""
            paragraph = cells[index].paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            _style_run(paragraph.add_run(text))
    document.add_paragraph()


def _build_docx(doc_result, out_path: Path) -> None:
    """Writes Docling's parsed items into a .docx with full-width native tables."""
    document = Document()
    style = document.styles["Normal"]
    style.font.name = LATIN_FONT
    style.font.size = Pt(10.5)

    doc = doc_result.document
    wrote_any = False

    for item, _level in doc.iterate_items():
        kind = type(item).__name__.lower()

        if "table" in kind and hasattr(item, "export_to_dataframe"):
            frame = item.export_to_dataframe()
            rows = [[str(c) for c in frame.columns]]
            rows += [[("" if v is None else str(v)) for v in record] for record in frame.values]
            _add_table(document, rows)
            wrote_any = True
            continue

        text = (getattr(item, "text", "") or "").strip()
        if not text:
            continue
        heading = "section" in kind or "title" in kind
        paragraph = document.add_paragraph()
        _style_run(paragraph.add_run(text), size_pt=13 if heading else 10.5, bold=heading)
        wrote_any = True

    if not wrote_any:
        raise HTTPException(
            status_code=422,
            detail="No readable content found — the scan resolution is too low. Re-scan at 300 dpi.",
        )

    document.save(out_path)


@app.post("/convert")
async def convert_pdf(
    file: UploadFile = File(...),
    ocr_enabled: bool = Form(True),
    preserve_layout: bool = Form(True),
    language_pack: str = Form("en_hi_mr"),
    ocr_language: str | None = Form(None),
):
    if not (file.filename or "").lower().endswith(".pdf"):
        raise HTTPException(status_code=415, detail="Only PDF uploads are supported.")

    # Imported lazily so the container boots fast and errors stay readable.
    from docling.datamodel.base_models import InputFormat
    from docling.datamodel.pipeline_options import PdfPipelineOptions
    from docling.document_converter import DocumentConverter, PdfFormatOption

    langs = {
        "en": ["eng"],
        "en_hi_mr": ["eng", "hin", "mar"],
        "en_es": ["eng", "spa"],
    }.get(language_pack, ["eng"])

    options = PdfPipelineOptions()
    options.do_ocr = ocr_enabled
    options.do_table_structure = preserve_layout
    if preserve_layout:
        options.table_structure_options.do_cell_matching = True
    if ocr_enabled:
        options.ocr_options.lang = langs

    converter = DocumentConverter(
        format_options={InputFormat.PDF: PdfFormatOption(pipeline_options=options)}
    )

    workdir = Path(tempfile.mkdtemp())
    pdf_path = workdir / (file.filename or "input.pdf")
    pdf_path.write_bytes(await file.read())

    try:
        result = converter.convert(str(pdf_path))
    except Exception as error:  # noqa: BLE001 - surfaced to the client verbatim
        raise HTTPException(status_code=422, detail=f"Could not parse this PDF: {error}") from error

    out_path = workdir / f"{pdf_path.stem}.docx"
    _build_docx(result, out_path)

    return FileResponse(
        out_path,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=out_path.name,
    )
